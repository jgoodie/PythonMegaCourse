#!/usr/local/opt/python3/bin/python3

import os
from re import split
from docx import Document

docdir = '/Users/jgoodman/Documents/HCP Anywhere/workspaces/PythonMegaCourse/docx2txt/docs'
txtdir = '/Users/jgoodman/Documents/HCP Anywhere/workspaces/PythonMegaCourse/docx2txt/txt'
os.chdir(docdir)
files = os.listdir(os.getcwd())

for x in files:
    os.chdir(docdir)
    doc = Document(x)
    #for p in doc.paragraphs:
     #   print(p)

# print(docdir + '/' + '20160629-EMEA-UK-LONDON-LONDON-RETAIL-HyperConverged-Marks and Spencer-Interview.docx')


# doc = Document(docdir + '/' + '20160629-EMEA-UK-LONDON-LONDON-RETAIL-HyperConverged-Marks and Spencer-Interview.docx')
# fulltext = []
# for x in doc.paragraphs:
#     fulltext.append(x.text)
# # return '\n'.join(fulltext)    

# for x in doc.paragraphs:
#     print(x.text)
    
    
    
    